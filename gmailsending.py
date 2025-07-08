import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
import os
import smtplib
from email.message import EmailMessage
from fpdf import FPDF
from datetime import datetime

# Paths and email config (adjust as needed)
cert_csv_path = r'D:\certificate\data.csv'
attendance_csv_path = r'D:\certificate\employee_attendance_dataset (1(3))_corrected.csv'
template_path = r'D:\certificate\ORG.docx'
docx_folder = r'D:\certificate\docxfolder'
pdf_folder = r'D:\certificate\pdf_folder'
employee_folder = r'D:\certificate\employee_folder'

YOUR_EMAIL = "sulaimansulaiman15795@gmail.com"
YOUR_PASSWORD = "ejjt myaf ddvg hdcn"

for folder in [docx_folder, pdf_folder, employee_folder]:
    os.makedirs(folder, exist_ok=True)

def send_email(to_email, subject, body, attachments=[]):
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = YOUR_EMAIL
    msg["To"] = to_email
    msg.set_content(body)

    for file_path in attachments:
        with open(file_path, "rb") as f:
            file_data = f.read()
            file_name = os.path.basename(file_path)
            msg.add_attachment(file_data, maintype="application", subtype="pdf", filename=file_name)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(YOUR_EMAIL, YOUR_PASSWORD)
        smtp.send_message(msg)

def personalize_paragraphs(doc, replacements):
    for p in doc.paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        if any(k in full_text for k in replacements):
            for i in range(len(p.runs)):
                p.runs[i].text = ''
            for key, val in replacements.items():
                full_text = full_text.replace(key, val)
            run = p.add_run(full_text)
            run.font.color.rgb = RGBColor(0, 0, 0)

def personalize_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = ''.join(run.text for run in p.runs)
                    if any(k in full_text for k in replacements):
                        for i in range(len(p.runs)):
                            p.runs[i].text = ''
                        for key, val in replacements.items():
                            full_text = full_text.replace(key, val)
                        run = p.add_run(full_text)
                        run.font.color.rgb = RGBColor(0, 0, 0)

class AttendancePDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, "Attendance Report", ln=True, align="C")
        self.ln(5)

print("Starting combined certificate and attendance generation and mailing...")

cert_df = pd.read_csv(cert_csv_path)
attendance_df = pd.read_csv(attendance_csv_path, encoding='utf-8')
attendance_df['Date'] = pd.to_datetime(attendance_df['Date'], errors='coerce')

for _, cert_row in cert_df.iterrows():
    try:
        # Extract cert data
        name = cert_row['name']
        reg_no = cert_row['reg_no']
        university = cert_row['university']
        course = cert_row['course']
        start_date = cert_row['start_date']
        end_date = cert_row['end_date']
        email = cert_row['email']
        issue_date = datetime.today().strftime('%d-%m-%Y')
        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_')).strip()

        # Create certificate PDF
        replacements = {
            '{name}': str(name),
            '{reg_no}': str(reg_no),
            '{university}': str(university),
            '{course}': str(course),
            '{start_date}': str(start_date),
            '{end_date}': str(end_date),
            '{issue_date}': issue_date
        }

        doc = Document(template_path)
        personalize_paragraphs(doc, replacements)
        personalize_tables(doc, replacements)
        docx_path = os.path.join(docx_folder, f"{safe_name}_certificate.docx")
        doc.save(docx_path)
        pdf_cert_path = os.path.join(pdf_folder, f"{safe_name}_certificate.pdf")
        convert(docx_path, pdf_cert_path)

        # Find attendance for this person by email
        attendance_person = attendance_df[attendance_df['email'] == email]
        if attendance_person.empty:
            print(f"⚠️ No attendance data for {name} ({email}), skipping attendance PDF.")
            attachments = [pdf_cert_path]
        else:
            emp_id = attendance_person['EmployeeID'].iloc[0]
            start_att_date = attendance_person['Date'].min().strftime('%d-%m-%Y')
            end_att_date = attendance_person['Date'].max().strftime('%d-%m-%Y')
            present_days = attendance_person[attendance_person['Status'].str.lower() == 'present'].shape[0]
            total_days = attendance_person.shape[0]
            percentage = (present_days / total_days) * 100 if total_days > 0 else 0

            emp_folder = os.path.join(employee_folder, f"{emp_id}_{safe_name}")
            os.makedirs(emp_folder, exist_ok=True)

            pdf_att = AttendancePDF()
            pdf_att.add_page()
            pdf_att.set_font("Arial", size=12)
            pdf_att.multi_cell(0, 10, f"This is to certify that {name} (Employee ID: {emp_id})")
            pdf_att.multi_cell(0, 10, f"has successfully attended work from {start_att_date} to {end_att_date}.")
            pdf_att.multi_cell(0, 10, f"The employee was present for {present_days} out of {total_days} working days.")
            pdf_att.multi_cell(0, 10, f"Attendance Percentage: {percentage:.2f}%")
            pdf_att.ln(10)
            pdf_att.set_font("Arial", "B", 12)
            pdf_att.cell(0, 10, f"Daily Attendance Sheet: {name}", ln=True)
            pdf_att.ln(5)
            pdf_att.set_font("Arial", "B", 11)
            pdf_att.cell(30, 10, "Date", border=1, align='C')
            pdf_att.cell(20, 10, "Present", border=1, align='C')
            pdf_att.cell(23, 10, "Absent", border=1, ln=True, align='C')
            pdf_att.set_font("Arial", "", 11)
            for _, r in attendance_person.iterrows():
                date_str = r['Date'].strftime('%d-%m-%Y')
                status = r['Status'].strip().lower()
                pdf_att.cell(30, 10, date_str, border=1)
                if status == "present":
                    pdf_att.cell(20, 10, "P", border=1, align='C')
                    pdf_att.cell(23, 10, "", border=1, ln=True)
                else:
                    pdf_att.cell(20, 10, "", border=1)
                    pdf_att.cell(23, 10, "A", border=1, ln=True)
            pdf_att_path = os.path.join(emp_folder, "attendance_report.pdf")
            pdf_att.output(pdf_att_path)

            attachments = [pdf_cert_path, pdf_att_path]

        # Email both certificate and attendance PDFs
        subject = "Your Internship Certificate and Attendance Report"
        body = f"Dear {name},\n\nPlease find attached your internship certificate and attendance report.\n\nBest regards,\nVDart Academy"

        send_email(email, subject, body, attachments)

        print(f"✅ Sent certificate and attendance to {name} ({email})")

    except Exception as e:
        print(f"❌ Error processing {name}: {e}")

print("All done!")